#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
core/parsing.py — ETAPA 1 do pipeline: LlamaParse.

Responsavel por transformar documento complexo (PDF escaneado, matricula, dossie
MentoRH, contrato, certidao, processo digitalizado) em Markdown limpo e navegavel.

Fluxo da API LlamaCloud:
  POST {base}/parsing/upload              -> {"id": job_id}
  GET  {base}/parsing/job/{id}            -> {"status": "SUCCESS"|"PENDING"|"ERROR"}
  GET  {base}/parsing/job/{id}/result/markdown -> {"markdown": "..."}

Fallback local: se nao houver chave LlamaParse cadastrada, cai para pdfplumber /
python-docx. O fallback e declarado no resultado para o operador saber a procedencia.
"""

from __future__ import annotations

import io
import time
from dataclasses import dataclass, field
from pathlib import Path

import requests

from .keys import KeyPool, PROVIDERS

TIMEOUT_UPLOAD = 180
TIMEOUT_POLL = 60
POLL_INTERVAL = 3
POLL_MAX_SEG = 600


@dataclass
class ParseResult:
    nome: str
    markdown: str
    origem: str                       # "llamaparse" | "pdfplumber" | "python-docx" | "texto"
    paginas: int = 0
    erro: str = ""
    metadados: dict = field(default_factory=dict)

    @property
    def ok(self) -> bool:
        return bool(self.markdown.strip()) and not self.erro


# --------------------------------------------------------------------------- #
# LLAMAPARSE
# --------------------------------------------------------------------------- #

def _parse_llamaparse(pool: KeyPool, nome: str, conteudo: bytes,
                      instrucao: str = "", modo: str = "balanced",
                      idioma: str = "pt") -> ParseResult | None:
    """Tenta o LlamaParse percorrendo todas as chaves ativas. None se nenhuma servir."""
    base = PROVIDERS["llamaparse"]["base"]
    tentadas: set[str] = set()

    while True:
        entry = pool.next_key("llamaparse")
        if entry is None or entry["key"] in tentadas:
            return None
        tentadas.add(entry["key"])
        token = entry["key"]
        headers = {"Authorization": f"Bearer {token}", "accept": "application/json"}

        data = {
            "language": idioma,
            "parse_mode": modo,
            "result_type": "markdown",
            "disable_ocr": "false",
            "skip_diagonal_text": "true",
            "do_not_unroll_columns": "false",
        }
        if instrucao:
            # instrucao em linguagem natural sobre o que extrair do documento
            data["parsing_instruction"] = instrucao

        try:
            up = requests.post(f"{base}/parsing/upload", headers=headers,
                               files={"file": (nome, io.BytesIO(conteudo))},
                               data=data, timeout=TIMEOUT_UPLOAD)
        except requests.RequestException as exc:
            return ParseResult(nome, "", "llamaparse", erro=f"rede: {exc}")

        if up.status_code in (429, 402):
            pool.marcar_cota("llamaparse", token)
            continue
        if up.status_code in (401, 403):
            pool.marcar_invalida("llamaparse", token, f"HTTP {up.status_code}")
            continue
        if up.status_code not in (200, 201):
            return ParseResult(nome, "", "llamaparse",
                               erro=f"upload HTTP {up.status_code}: {up.text[:200]}")

        job_id = up.json().get("id")
        if not job_id:
            return ParseResult(nome, "", "llamaparse", erro="resposta sem job id")

        # polling
        inicio = time.time()
        status = "PENDING"
        while time.time() - inicio < POLL_MAX_SEG:
            time.sleep(POLL_INTERVAL)
            try:
                st = requests.get(f"{base}/parsing/job/{job_id}",
                                  headers=headers, timeout=TIMEOUT_POLL)
            except requests.RequestException:
                continue
            if st.status_code != 200:
                continue
            status = (st.json().get("status") or "").upper()
            if status in ("SUCCESS", "COMPLETED", "ERROR", "FAILED", "CANCELED"):
                break

        if status not in ("SUCCESS", "COMPLETED"):
            return ParseResult(nome, "", "llamaparse",
                               erro=f"job {job_id} terminou como {status or 'TIMEOUT'}")

        try:
            res = requests.get(f"{base}/parsing/job/{job_id}/result/markdown",
                               headers=headers, timeout=TIMEOUT_POLL)
        except requests.RequestException as exc:
            return ParseResult(nome, "", "llamaparse", erro=f"resultado: {exc}")

        if res.status_code != 200:
            return ParseResult(nome, "", "llamaparse",
                               erro=f"resultado HTTP {res.status_code}")

        payload = res.json()
        md = payload.get("markdown") or payload.get("text") or ""
        pool.marcar_uso("llamaparse", token)
        return ParseResult(nome, md, "llamaparse",
                           paginas=int(payload.get("job_metadata", {})
                                       .get("job_pages", 0) or 0),
                           metadados={"job_id": job_id})


# --------------------------------------------------------------------------- #
# FALLBACK LOCAL
# --------------------------------------------------------------------------- #

def _parse_local(nome: str, conteudo: bytes) -> ParseResult:
    baixo = nome.lower()

    if baixo.endswith((".txt", ".md")):
        return ParseResult(nome, conteudo.decode("utf-8", errors="replace"), "texto")

    if baixo.endswith(".pdf"):
        try:
            import pdfplumber
            partes, n = [], 0
            with pdfplumber.open(io.BytesIO(conteudo)) as pdf:
                for i, pagina in enumerate(pdf.pages, 1):
                    n = i
                    txt = pagina.extract_text() or ""
                    partes.append(f"\n\n<!-- pagina {i} -->\n{txt}")
                    for tabela in (pagina.extract_tables() or []):
                        linhas = ["| " + " | ".join((c or "") for c in row) + " |"
                                  for row in tabela if row]
                        if linhas:
                            partes.append("\n" + "\n".join(linhas))
            texto = "".join(partes).strip()
            if not texto:
                return ParseResult(nome, "", "pdfplumber", paginas=n,
                                   erro="PDF sem camada de texto — precisa de OCR "
                                        "(cadastre uma chave LlamaParse)")
            return ParseResult(nome, texto, "pdfplumber", paginas=n)
        except Exception as exc:
            return ParseResult(nome, "", "pdfplumber", erro=str(exc))

    if baixo.endswith(".docx"):
        try:
            import docx
            d = docx.Document(io.BytesIO(conteudo))
            partes = [p.text for p in d.paragraphs if p.text.strip()]
            for t in d.tables:
                for row in t.rows:
                    partes.append("| " + " | ".join(c.text.strip() for c in row.cells) + " |")
            return ParseResult(nome, "\n".join(partes), "python-docx")
        except Exception as exc:
            return ParseResult(nome, "", "python-docx", erro=str(exc))

    if baixo.endswith(".odt"):
        try:
            import zipfile, re, html
            with zipfile.ZipFile(io.BytesIO(conteudo)) as z:
                xml = z.read("content.xml").decode("utf-8")
            xml = re.sub(r"</text:p>", "\n", xml)
            xml = re.sub(r"<[^>]+>", "", xml)
            return ParseResult(nome, html.unescape(xml).strip(), "odt")
        except Exception as exc:
            return ParseResult(nome, "", "odt", erro=str(exc))

    return ParseResult(nome, "", "desconhecido",
                       erro=f"formato nao suportado localmente: {nome}")


# --------------------------------------------------------------------------- #
# API PUBLICA
# --------------------------------------------------------------------------- #

INSTRUCAO_PADRAO = (
    "Este e um documento administrativo do Poder Judiciario do Maranhao. "
    "Preserve integralmente: nomes completos, matriculas, numeros de processo, "
    "numeros e anos de portaria, resolucao, oficio, decisao e manifestacao, "
    "cargos, unidades de lotacao, datas e valores monetarios. "
    "Converta tabelas em tabelas Markdown sem omitir linhas. "
    "Nao resuma, nao interprete e nao corrija o conteudo."
)


def parse_documento(pool: KeyPool, nome: str, conteudo: bytes,
                    instrucao: str = INSTRUCAO_PADRAO,
                    preferir_llamaparse: bool = True) -> ParseResult:
    """Extrai um documento. LlamaParse primeiro; fallback local se indisponivel."""
    if preferir_llamaparse and pool.disponiveis("llamaparse"):
        r = _parse_llamaparse(pool, nome, conteudo, instrucao)
        if r is not None and r.ok:
            return r
        if r is not None and r.erro:
            local = _parse_local(nome, conteudo)
            if local.ok:
                local.erro = f"(LlamaParse falhou: {r.erro} — usado fallback local)"
                return local
            return r
    return _parse_local(nome, conteudo)


def parse_lote(pool: KeyPool, arquivos: list[tuple[str, bytes]],
               instrucao: str = INSTRUCAO_PADRAO) -> list[ParseResult]:
    return [parse_documento(pool, nome, blob, instrucao) for nome, blob in arquivos]


def consolidar(resultados: list[ParseResult], limite_chars: int = 120_000) -> str:
    """Concatena os markdowns extraidos, com cabecalho de procedencia por anexo."""
    blocos = []
    for r in resultados:
        cab = (f"===== ANEXO: {r.nome} | extracao: {r.origem}"
               f"{f' | {r.paginas} pag.' if r.paginas else ''} =====")
        corpo = r.markdown if r.ok else f"[FALHA NA EXTRACAO: {r.erro}]"
        blocos.append(f"{cab}\n{corpo}")
    texto = "\n\n".join(blocos)
    if len(texto) > limite_chars:
        texto = texto[:limite_chars] + "\n\n[TRUNCADO POR LIMITE DE CONTEXTO]"
    return texto
