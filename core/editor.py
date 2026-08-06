#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
core/editor.py — Editor de texto tipo Word dentro da aplicacao.

Pilha escolhida (bibliotecas Python ja existentes para este fim):
  streamlit-quill  -> editor WYSIWYG (Quill.js) com toolbar de negrito, italico,
                      sublinhado, alinhamento, recuo, listas, cabecalhos e cores.
                      Retorna HTML, que e o formato nativo de trabalho aqui.
  python-docx      -> conversao HTML -> .docx preservando os atributos de formatacao.

Este modulo faz:
  - texto plano  -> HTML inicial (heuristica de titulo/paragrafo/tabela);
  - HTML         -> texto plano (para busca e para o FTS do SQLite);
  - HTML         -> .docx (negrito, italico, sublinhado, alinhamento, recuo,
                    espacamento entre linhas, listas e tabelas);
  - CSS da folha A4, para o editor parecer uma pagina do Word.
"""

from __future__ import annotations

import html as htmllib
import re
from html.parser import HTMLParser
from pathlib import Path

# --------------------------------------------------------------------------- #
# CONFIGURACAO DA PAGINA (padrao de peca administrativa)
# --------------------------------------------------------------------------- #

PAGINA = {
    "fonte": "Times New Roman",
    "tamanho_pt": 12,
    "margem_sup_cm": 3.0,
    "margem_inf_cm": 2.0,
    "margem_esq_cm": 3.0,
    "margem_dir_cm": 2.0,
    "entrelinhas": 1.5,
    "espaco_depois_pt": 6,
    "recuo_primeira_linha_cm": 1.25,
}

TOOLBAR_QUILL = [
    [{"header": [1, 2, 3, False]}],
    ["bold", "italic", "underline", "strike"],
    [{"align": []}],
    [{"list": "ordered"}, {"list": "bullet"}],
    [{"indent": "-1"}, {"indent": "+1"}],
    [{"color": []}, {"background": []}],
    ["blockquote"],
    ["clean"],
]


def css_folha(escuro: bool = True) -> str:
    """CSS que transforma o container do editor numa folha A4."""
    papel = "#ffffff"
    tinta = "#111111"
    fundo = "#2b2b2b" if escuro else "#e9e9e9"
    return f"""
<style>
.folha-wrap {{
    background: {fundo};
    padding: 28px 0 34px 0;
    border-radius: 6px;
    display: flex;
    justify-content: center;
}}
.folha {{
    background: {papel};
    color: {tinta};
    width: 21cm;
    min-height: 12cm;
    padding: {PAGINA['margem_sup_cm']}cm {PAGINA['margem_dir_cm']}cm
             {PAGINA['margem_inf_cm']}cm {PAGINA['margem_esq_cm']}cm;
    box-shadow: 0 4px 22px rgba(0,0,0,.45);
    font-family: "{PAGINA['fonte']}", Georgia, serif;
    font-size: {PAGINA['tamanho_pt']}pt;
    line-height: {PAGINA['entrelinhas']};
    text-align: justify;
    box-sizing: border-box;
}}
.folha p {{ margin: 0 0 {PAGINA['espaco_depois_pt']}pt 0; }}
.folha h1, .folha h2, .folha h3 {{
    text-align: center; font-weight: 700; margin: 14pt 0 10pt 0;
    font-size: {PAGINA['tamanho_pt']}pt; text-transform: uppercase;
}}
.folha table {{ border-collapse: collapse; width: 100%; margin: 8pt 0; }}
.folha td, .folha th {{ border: 1px solid #444; padding: 4pt 6pt; font-size: 11pt; }}
.folha .cabecalho {{
    text-align: center; border-bottom: 1px solid #999;
    padding-bottom: 8pt; margin-bottom: 16pt; font-size: 11pt;
}}
.folha .rodape {{
    border-top: 1px solid #999; margin-top: 20pt; padding-top: 6pt;
    font-size: 9pt; color: #555; text-align: center;
}}
.folha mark.faltante {{
    background: #ffe08a; color: #7a4b00; padding: 0 3px;
    border-radius: 3px; font-weight: 600;
}}
/* editor Quill sobre fundo claro, para leitura fiel */
.stQuill, .stQuill .ql-container {{ background: #ffffff !important; }}
.stQuill .ql-editor {{
    font-family: "{PAGINA['fonte']}", Georgia, serif !important;
    font-size: {PAGINA['tamanho_pt']}pt !important;
    line-height: {PAGINA['entrelinhas']} !important;
    color: #111 !important;
    min-height: 460px;
    text-align: justify;
}}
.stQuill .ql-toolbar {{ background: #f2f2f2 !important; border-radius: 6px 6px 0 0; }}
</style>
"""


# --------------------------------------------------------------------------- #
# TEXTO PLANO -> HTML
# --------------------------------------------------------------------------- #

TITULOS = ("INFORMAÇÃO", "INFORMACAO", "DESPACHO", "OFÍCIO", "OFICIO",
           "MEMORANDO", "MANIFESTAÇÃO", "MANIFESTACAO", "REQUERIMENTO")

RODAPE_TEC = re.compile(r"^(Tipo classificado|Modelo-base utilizado|Norma invocada|"
                        r"Semestre de referência|Semestre de referencia|"
                        r"Campos faltantes)\s*:", re.IGNORECASE)

FALTANTE = re.compile(r"\[DADO FALTANTE:\s*([^\]]+)\]")


def _linha_tabela(l: str) -> bool:
    return l.strip().startswith("|") and l.strip().endswith("|")


def texto_para_html(texto: str, marcar_faltantes: bool = True) -> str:
    """Converte a saida do pipeline em HTML editavel, inferindo estrutura."""
    linhas = (texto or "").split("\n")
    out: list[str] = []
    buffer_tabela: list[str] = []

    def descarrega_tabela() -> None:
        nonlocal buffer_tabela
        if not buffer_tabela:
            return
        linhas_uteis = [l for l in buffer_tabela
                        if not re.fullmatch(r"\|[\s\-\|:]+\|", l.strip())]
        if linhas_uteis:
            out.append("<table>")
            for i, l in enumerate(linhas_uteis):
                celulas = [c.strip() for c in l.strip().strip("|").split("|")]
                tag = "th" if i == 0 else "td"
                out.append("<tr>" + "".join(
                    f"<{tag}>{htmllib.escape(c)}</{tag}>" for c in celulas) + "</tr>")
            out.append("</table>")
        buffer_tabela = []

    for linha in linhas:
        bruto = linha.strip()

        if _linha_tabela(bruto):
            buffer_tabela.append(bruto)
            continue
        descarrega_tabela()

        if not bruto:
            continue

        seguro = htmllib.escape(bruto)
        if marcar_faltantes:
            seguro = FALTANTE.sub(
                r'<mark class="faltante">[DADO FALTANTE: \1]</mark>', seguro)

        if bruto.upper() in TITULOS or bruto.rstrip(".").upper() in TITULOS:
            out.append(f"<h2>{seguro}</h2>")
        elif RODAPE_TEC.match(bruto):
            out.append(f'<p style="font-size:10pt;color:#555">{seguro}</p>')
        elif bruto == bruto.upper() and 3 < len(bruto) < 70 and not bruto[0].isdigit():
            out.append(f'<p style="text-align:center"><strong>{seguro}</strong></p>')
        elif re.match(r"^[a-z]\)\s", bruto) or bruto.startswith("- "):
            out.append(f'<p style="margin-left:1.25cm">{seguro}</p>')
        else:
            out.append(f"<p>{seguro}</p>")

    descarrega_tabela()
    return "\n".join(out) or "<p></p>"


# --------------------------------------------------------------------------- #
# HTML -> TEXTO PLANO
# --------------------------------------------------------------------------- #

class _Plano(HTMLParser):
    def __init__(self):
        super().__init__()
        self.partes: list[str] = []
        self._celula = False

    def handle_starttag(self, tag, attrs):
        if tag in ("br",):
            self.partes.append("\n")
        elif tag in ("td", "th"):
            self._celula = True
            self.partes.append(" | ")

    def handle_endtag(self, tag):
        if tag in ("p", "div", "h1", "h2", "h3", "li", "tr", "table", "blockquote"):
            self.partes.append("\n")
        elif tag in ("td", "th"):
            self._celula = False

    def handle_data(self, data):
        self.partes.append(data)


def html_para_texto(html: str) -> str:
    p = _Plano()
    p.feed(html or "")
    texto = "".join(p.partes)
    texto = htmllib.unescape(texto)
    texto = re.sub(r"[ \t]+", " ", texto)
    texto = re.sub(r"\n{3,}", "\n\n", texto)
    return "\n".join(l.strip(" |").strip() for l in texto.split("\n")).strip()


# --------------------------------------------------------------------------- #
# HTML -> DOCX
# --------------------------------------------------------------------------- #

class _ParaDocx(HTMLParser):
    """Percorre o HTML do Quill e monta o .docx preservando a formatacao."""

    BLOCOS = {"p", "h1", "h2", "h3", "li", "blockquote"}

    def __init__(self, doc):
        super().__init__()
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        self.WD = WD_ALIGN_PARAGRAPH
        self.doc = doc
        self.par = None
        self.negrito = self.italico = self.sublinhado = self.tachado = False
        self.alinhamento = None
        self.recuo = 0
        self.titulo = False
        self.lista = None
        self.tabela = None
        self.linha_celulas: list[str] = []
        self.tabela_linhas: list[list[str]] = []
        self.em_celula = False
        self.buffer_celula = ""

    # -------- utilidades --------

    @staticmethod
    def _attrs(attrs) -> dict:
        return {k: (v or "") for k, v in attrs}

    def _alinhamento_de(self, style: str, classe: str):
        alvo = f"{style} {classe}"
        if "center" in alvo:
            return self.WD.CENTER
        if "right" in alvo:
            return self.WD.RIGHT
        if "justify" in alvo:
            return self.WD.JUSTIFY
        if "left" in alvo:
            return self.WD.LEFT
        return None

    @staticmethod
    def _recuo_de(style: str, classe: str) -> float:
        m = re.search(r"margin-left:\s*([\d.]+)cm", style)
        if m:
            return float(m.group(1))
        m = re.search(r"ql-indent-(\d+)", classe)
        if m:
            return 1.25 * int(m.group(1))
        return 0.0

    def _novo_par(self):
        from docx.shared import Pt, Cm
        estilo = "List Bullet" if self.lista == "ul" else (
            "List Number" if self.lista == "ol" else None)
        self.par = self.doc.add_paragraph(style=estilo) if estilo \
            else self.doc.add_paragraph()
        pf = self.par.paragraph_format
        pf.line_spacing = PAGINA["entrelinhas"]
        pf.space_after = Pt(PAGINA["espaco_depois_pt"])
        pf.alignment = self.alinhamento if self.alinhamento is not None else (
            self.WD.CENTER if self.titulo else self.WD.JUSTIFY)
        if self.recuo:
            pf.left_indent = Cm(self.recuo)

    # -------- parser --------

    def handle_starttag(self, tag, attrs):
        a = self._attrs(attrs)
        style, classe = a.get("style", ""), a.get("class", "")

        if tag == "table":
            self.tabela_linhas = []
            self.tabela = True
            return
        if tag == "tr":
            self.linha_celulas = []
            return
        if tag in ("td", "th"):
            self.em_celula = True
            self.buffer_celula = ""
            return

        if tag in ("ul", "ol"):
            self.lista = tag
            return

        if tag in self.BLOCOS:
            self.titulo = tag in ("h1", "h2", "h3")
            self.alinhamento = self._alinhamento_de(style, classe)
            self.recuo = self._recuo_de(style, classe)
            self._novo_par()
            if self.titulo:
                self.negrito = True
            return

        if tag in ("strong", "b"):
            self.negrito = True
        elif tag in ("em", "i"):
            self.italico = True
        elif tag in ("u", "ins"):
            self.sublinhado = True
        elif tag in ("s", "del", "strike"):
            self.tachado = True
        elif tag == "br" and self.par is not None:
            self.par.add_run().add_break()
        elif tag == "mark":
            self.negrito = True

    def handle_endtag(self, tag):
        if tag in ("td", "th"):
            self.linha_celulas.append(self.buffer_celula.strip())
            self.em_celula = False
            return
        if tag == "tr":
            if self.linha_celulas:
                self.tabela_linhas.append(self.linha_celulas)
            return
        if tag == "table":
            self._emitir_tabela()
            self.tabela = False
            return
        if tag in ("ul", "ol"):
            self.lista = None
            return
        if tag in self.BLOCOS:
            if self.titulo:
                self.negrito = False
                self.titulo = False
            self.par = None
            self.alinhamento = None
            self.recuo = 0
            return
        if tag in ("strong", "b", "mark"):
            self.negrito = False
        elif tag in ("em", "i"):
            self.italico = False
        elif tag in ("u", "ins"):
            self.sublinhado = False
        elif tag in ("s", "del", "strike"):
            self.tachado = False

    def handle_data(self, data):
        if self.em_celula:
            self.buffer_celula += data
            return
        if not data.strip():
            return
        if self.par is None:
            self._novo_par()
        run = self.par.add_run(htmllib.unescape(data))
        run.bold = self.negrito
        run.italic = self.italico
        run.underline = self.sublinhado
        if self.tachado:
            run.font.strike = True

    def _emitir_tabela(self):
        if not self.tabela_linhas:
            return
        from docx.shared import Pt
        colunas = max(len(l) for l in self.tabela_linhas)
        t = self.doc.add_table(rows=0, cols=colunas)
        t.style = "Table Grid"
        for i, linha in enumerate(self.tabela_linhas):
            celulas = t.add_row().cells
            for j in range(colunas):
                texto = linha[j] if j < len(linha) else ""
                celulas[j].text = texto
                for p in celulas[j].paragraphs:
                    for r in p.runs:
                        r.font.size = Pt(11)
                        if i == 0:
                            r.bold = True
        self.doc.add_paragraph()


def html_para_docx(corpo_html: str, destino: Path,
                   cabecalho_html: str = "", rodape_html: str = "") -> Path:
    """Gera o .docx a partir do HTML do editor, com margens de peca administrativa."""
    from docx import Document
    from docx.shared import Pt, Cm

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(PAGINA["margem_sup_cm"])
    sec.bottom_margin = Cm(PAGINA["margem_inf_cm"])
    sec.left_margin = Cm(PAGINA["margem_esq_cm"])
    sec.right_margin = Cm(PAGINA["margem_dir_cm"])

    normal = doc.styles["Normal"]
    normal.font.name = PAGINA["fonte"]
    normal.font.size = Pt(PAGINA["tamanho_pt"])

    if cabecalho_html:
        cab = sec.header.paragraphs[0]
        cab.text = html_para_texto(cabecalho_html)
        for r in cab.runs:
            r.font.size = Pt(10)
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        cab.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if rodape_html:
        rod = sec.footer.paragraphs[0]
        rod.text = html_para_texto(rodape_html)
        for r in rod.runs:
            r.font.size = Pt(8)
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        rod.alignment = WD_ALIGN_PARAGRAPH.CENTER

    parser = _ParaDocx(doc)
    parser.feed(corpo_html or "")

    destino = Path(destino)
    destino.parent.mkdir(parents=True, exist_ok=True)
    doc.save(destino)
    return destino


# --------------------------------------------------------------------------- #
# PRE-VISUALIZACAO
# --------------------------------------------------------------------------- #

def folha_html(corpo_html: str, cabecalho_html: str = "",
               rodape_html: str = "") -> str:
    """Monta a folha A4 completa para pre-visualizacao."""
    partes = ['<div class="folha-wrap"><div class="folha">']
    if cabecalho_html:
        partes.append(f'<div class="cabecalho">{cabecalho_html}</div>')
    partes.append(corpo_html or "<p></p>")
    if rodape_html:
        partes.append(f'<div class="rodape">{rodape_html}</div>')
    partes.append("</div></div>")
    return "".join(partes)


CABECALHO_PADRAO = (
    '<strong>PODER JUDICIÁRIO DO ESTADO DO MARANHÃO</strong><br>'
    'TRIBUNAL DE JUSTIÇA<br>'
    'Coordenadoria de Acompanhamento e Desenvolvimento na Carreira — CAEDNC'
)

RODAPE_PADRAO = (
    'Documento gerado pelo Redator COCARREIRA/CAEDNC — '
    'conferir antes da expedição no DigiDoc.'
)
